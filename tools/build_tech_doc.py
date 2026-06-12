#!/usr/bin/env python
"""Render TECHNICAL_REQUIREMENTS.md into a professionally formatted Word document.

Builds a branded .docx (Northeastern red accents, cover page, auto table of
contents, styled headings, code blocks, and tables) from the markdown source.

Usage:
    /opt/anaconda3/bin/python tools/build_tech_doc.py
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "TECHNICAL_REQUIREMENTS.md")
OUT = os.path.join(ROOT, "Northeastern_Innovation_Dashboard_Technical_Requirements.docx")

# Northeastern brand palette
NU_RED = RGBColor(0xC8, 0x10, 0x2E)
NU_DARK = RGBColor(0x20, 0x20, 0x20)
NU_GREY = RGBColor(0x5A, 0x5A, 0x5A)
NU_LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
CODE_BG = "F4F4F4"
HEADER_FILL = "C8102E"
ROW_ALT_FILL = "F7E9EC"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


# ----------------------------------------------------------------------------
# Low-level XML helpers
# ----------------------------------------------------------------------------
def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_bottom_border(paragraph, color="C8102E", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_break(doc):
    doc.add_page_break()


def add_toc_field(doc):
    """Insert a Word TOC field that auto-populates on open/update."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


# ----------------------------------------------------------------------------
# Inline markdown (bold, code, links) -> runs
# ----------------------------------------------------------------------------
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline_runs(paragraph, text, base_size=10.5, base_color=NU_DARK):
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = base_color
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = NU_RED
        elif token.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            label = m.group(1) if m else token
            run = paragraph.add_run(label)
            run.font.name = MONO_FONT
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = NU_GREY
        else:
            run = paragraph.add_run(token)
            run.font.color.rgb = base_color
        run.font.size = run.font.size or Pt(base_size)
        if not run.font.name:
            run.font.name = BODY_FONT


# ----------------------------------------------------------------------------
# Document styling
# ----------------------------------------------------------------------------
def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NU_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (16, NU_RED, True),
        "Heading 2": (13, NU_RED, True),
        "Heading 3": (11.5, NU_DARK, True),
    }
    for name, (size, color, bold) in heading_specs.items():
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Technical Requirements\n& Handoff Document")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = NU_RED
    run.font.name = BODY_FONT

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(rule, size="18")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Northeastern University\nInnovation Intelligence Dashboard")
    r.font.size = Pt(16)
    r.font.color.rgb = NU_DARK
    r.font.name = BODY_FONT

    for _ in range(8):
        doc.add_paragraph()

    meta = [
        ("Document type", "Technical handoff / system design reference"),
        ("Audience", "Incoming project manager and engineering team"),
        ("Status", "Current implementation as of June 2026"),
        ("Last updated", "2026-06-07"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(f"{label}:  ")
        lr.bold = True
        lr.font.size = Pt(10.5)
        lr.font.color.rgb = NU_RED
        lr.font.name = BODY_FONT
        vr = p.add_run(value)
        vr.font.size = Pt(10.5)
        vr.font.color.rgb = NU_GREY
        vr.font.name = BODY_FONT
    add_page_break(doc)


def add_code_block(doc, lines):
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, CODE_BG)
    run = p.add_run(text)
    run.font.name = MONO_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = NU_DARK


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(header):
        set_cell_background(hdr_cells[i], HEADER_FILL)
        para = hdr_cells[i].paragraphs[0]
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        run = para.add_run(htext.strip())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = BODY_FONT
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, ctext in enumerate(row):
            if c_idx >= len(cells):
                continue
            if r_idx % 2 == 1:
                set_cell_background(cells[c_idx], ROW_ALT_FILL)
            para = cells[c_idx].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            add_inline_runs(para, ctext.strip(), base_size=9.5)
    doc.add_paragraph()


# ----------------------------------------------------------------------------
# Markdown parsing
# ----------------------------------------------------------------------------
def parse_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(line):
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line)) and "-" in line


def render_body(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip the top H1 title block + leading metadata (already on cover page)
        if i < 12 and (stripped.startswith("# ") or stripped.startswith("**") or stripped == ""):
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block)
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < n and is_separator_row(lines[i + 1]):
            header = parse_table_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            # Start each numbered top-level section on context; keep flow tidy
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Blockquote (callout)
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, "FBF0F2")
            add_bottom_border(p, size="4")
            add_inline_runs(p, text, base_size=10)
            i += 1
            continue

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m_num.group(2), base_size=10.5)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:], base_size=10.5)
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped, base_size=10.5)
        i += 1


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Northeastern University Innovation Intelligence Dashboard — Technical Requirements")
    run.font.size = Pt(8)
    run.font.color.rgb = NU_GREY
    run.font.name = BODY_FONT


def main():
    with open(SRC, "r", encoding="utf-8") as fh:
        md_text = fh.read()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    configure_styles(doc)
    build_cover(doc)

    # Table of contents
    toc_heading = doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)
    add_page_break(doc)

    render_body(doc, md_text)
    add_footer(doc)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
